from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from io import BytesIO
import re

from docx import Document
from flask import Blueprint, jsonify, redirect, render_template, request, send_file, session, url_for
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill

from models import get_db, pool


school_bp = Blueprint("school", __name__, url_prefix="/school")


def _company_id():
    return session.get("company_id") if session.get("user_id") else None


def _text(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _number(value, default=0):
    try:
        return max(0, int(Decimal(_text(value).replace(" ", "").replace(",", "."))))
    except (InvalidOperation, ValueError, TypeError):
        return default


def _money(value, default=Decimal("0")):
    try:
        return max(Decimal("0"), Decimal(_text(value).replace(" ", "").replace(",", ".")))
    except (InvalidOperation, ValueError, TypeError):
        return default


def _class_name(value):
    name = re.sub(r"\s+", "", _text(value)).upper()
    # В Word параллель третьих классов записана кириллической «З» вместо цифры 3.
    if len(name) >= 2 and name[0] == "З" and not name[1].isdigit():
        name = "3" + name[1:]
    # В исходной таблице буква «З» один раз записана цифрой 3: «1 3».
    if len(name) == 2 and name[1] == "3":
        name = name[0] + "З"
    return name


def _class_sort(name):
    match = re.match(r"^(\d+)(.*)$", name or "")
    if not match:
        return 9999
    grade = int(match.group(1))
    letter = match.group(2) or ""
    return grade * 100 + (ord(letter[0]) if letter else 0)


def _ensure_class(cur, company_id, name):
    name = _class_name(name)
    if not name:
        return None
    cur.execute("""
        INSERT INTO school_classes(company_id, name, sort_order)
        VALUES (%s, %s, %s)
        ON CONFLICT(company_id, name)
        DO UPDATE SET is_active = TRUE
        RETURNING id
    """, (company_id, name, _class_sort(name)))
    return cur.fetchone()["id"]


def _prices_for_date(cur, company_id, meal_date):
    cur.execute("""
        SELECT free_price, paid_price, effective_from
        FROM school_meal_prices
        WHERE company_id = %s AND effective_from <= %s
        ORDER BY effective_from DESC LIMIT 1
    """, (company_id, meal_date))
    return cur.fetchone() or {"free_price": Decimal("0"), "paid_price": Decimal("0"), "effective_from": meal_date}


def _style_sheet(ws, widths):
    fill = PatternFill("solid", fgColor="312E81")
    for cell in ws[1]:
        cell.fill = fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for index, width in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + index)].width = width


@school_bp.route("")
def index():
    return redirect(url_for("school.meals"))


@school_bp.route("/leaders")
def leaders():
    company_id = _company_id()
    if not company_id:
        return redirect("/login")
    search = request.args.get("q", "").strip()
    conn = get_db()
    try:
        cur = conn.cursor()
        params = [company_id]
        where = "WHERE l.company_id = %s"
        if search:
            where += " AND (l.full_name ILIKE %s OR c.name ILIKE %s OR l.phone ILIKE %s)"
            like = f"%{search}%"
            params.extend([like, like, like])
        cur.execute(f"""
            SELECT l.id, l.full_name, l.room, l.phone, l.class_id, c.name AS class_name
            FROM school_class_leaders l
            LEFT JOIN school_classes c ON c.id = l.class_id
            {where}
            ORDER BY c.sort_order, c.name, l.full_name
        """, params)
        rows = cur.fetchall()
        cur.execute("SELECT id, name FROM school_classes WHERE company_id=%s AND is_active=TRUE ORDER BY sort_order, name", (company_id,))
        classes = cur.fetchall()
    finally:
        pool.putconn(conn)
    return render_template("school_leaders.html", leaders=rows, classes=classes, search=search)


@school_bp.route("/leaders/save", methods=["POST"])
def save_leader():
    company_id = _company_id()
    if not company_id:
        return redirect("/login")
    leader_id = request.form.get("id", type=int)
    full_name = request.form.get("full_name", "").strip()
    class_name = request.form.get("class_name", "").strip()
    room = request.form.get("room", "").strip()
    phone = request.form.get("phone", "").strip()
    if not full_name:
        return redirect(url_for("school.leaders", error="Укажите ФИО"))
    conn = get_db()
    try:
        cur = conn.cursor()
        class_id = _ensure_class(cur, company_id, class_name) if class_name else None
        if leader_id:
            cur.execute("""
                UPDATE school_class_leaders
                SET full_name=%s, class_id=%s, room=%s, phone=%s, updated_at=NOW()
                WHERE id=%s AND company_id=%s
            """, (full_name, class_id, room, phone, leader_id, company_id))
        else:
            existing = None
            if class_id:
                cur.execute("SELECT id FROM school_class_leaders WHERE company_id=%s AND class_id=%s", (company_id, class_id))
                existing = cur.fetchone()
            if existing:
                cur.execute("UPDATE school_class_leaders SET full_name=%s, room=%s, phone=%s, updated_at=NOW() WHERE id=%s", (full_name, room, phone, existing["id"]))
            else:
                cur.execute("""
                    INSERT INTO school_class_leaders(company_id, class_id, full_name, room, phone)
                    VALUES (%s,%s,%s,%s,%s)
                """, (company_id, class_id, full_name, room, phone))
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        pool.putconn(conn)
    return redirect(url_for("school.leaders", success="Сохранено"))


@school_bp.route("/leaders/<int:leader_id>/delete", methods=["POST"])
def delete_leader(leader_id):
    company_id = _company_id()
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM school_class_leaders WHERE id=%s AND company_id=%s", (leader_id, company_id))
        conn.commit()
    finally:
        pool.putconn(conn)
    return redirect(url_for("school.leaders", success="Удалено"))


@school_bp.route("/leaders/import", methods=["POST"])
def import_leaders():
    company_id = _company_id()
    upload = request.files.get("file")
    if not company_id or not upload or not upload.filename:
        return redirect(url_for("school.leaders", error="Выберите файл"))
    filename = upload.filename.lower()
    parsed = []
    try:
        if filename.endswith(".docx"):
            doc = Document(upload)
            if not doc.tables:
                raise ValueError("В Word-файле нет таблицы")
            for row in doc.tables[0].rows[1:]:
                cells = [_text(cell.text) for cell in row.cells]
                if len(cells) >= 5 and cells[1]:
                    parsed.append((cells[1], cells[2], cells[3], cells[4]))
        elif filename.endswith(".xlsx"):
            ws = load_workbook(upload, read_only=True, data_only=True).active
            headers = {_text(c.value).lower(): i for i, c in enumerate(ws[1])}
            def column(*names):
                return next((headers[n] for n in names if n in headers), None)
            name_col = column("фио", "аты-жөні", "классный руководитель", "ф.и.о.")
            class_col = column("класс", "сыныбы / тобы", "сынып")
            room_col = column("кабинет")
            phone_col = column("телефон", "ұялы телефоны")
            if name_col is None:
                # Поддержка приложенного файла, где классы расположены блоками без заголовков.
                class_names = []
                for row in ws.iter_rows(values_only=True):
                    for cell in row:
                        candidate = _class_name(cell)
                        if re.match(r"^(?:[0-9]|10|11)[А-ЯӘІҢҒҮҰҚӨҺ]$", candidate):
                            class_names.append(candidate)
                conn = get_db()
                try:
                    cur = conn.cursor()
                    for candidate in dict.fromkeys(class_names):
                        _ensure_class(cur, company_id, candidate)
                    conn.commit()
                finally:
                    pool.putconn(conn)
                return redirect(url_for("school.leaders", success=f"Загружено классов: {len(set(class_names))}"))
            else:
                for row in ws.iter_rows(min_row=2, values_only=True):
                    name = _text(row[name_col])
                    if name:
                        parsed.append((name, _text(row[class_col]) if class_col is not None else "", _text(row[room_col]) if room_col is not None else "", _text(row[phone_col]) if phone_col is not None else ""))
        else:
            raise ValueError("Поддерживаются .docx и .xlsx")
    except Exception as exc:
        return redirect(url_for("school.leaders", error=f"Ошибка импорта: {exc}"))

    conn = get_db()
    try:
        cur = conn.cursor()
        for name, class_name, room, phone in parsed:
            class_id = _ensure_class(cur, company_id, class_name) if class_name else None
            cur.execute("""
                SELECT id FROM school_class_leaders
                WHERE company_id=%s AND (class_id=%s OR (class_id IS NULL AND %s IS NULL))
                ORDER BY id LIMIT 1
            """, (company_id, class_id, class_id))
            existing = cur.fetchone()
            if existing:
                cur.execute("UPDATE school_class_leaders SET full_name=%s, room=%s, phone=%s, updated_at=NOW() WHERE id=%s", (name, room, phone, existing["id"]))
            else:
                cur.execute("INSERT INTO school_class_leaders(company_id,class_id,full_name,room,phone) VALUES(%s,%s,%s,%s,%s)", (company_id, class_id, name, room, phone))
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        pool.putconn(conn)
    return redirect(url_for("school.leaders", success=f"Загружено: {len(parsed)}"))


@school_bp.route("/leaders/export.xlsx")
def export_leaders():
    company_id = _company_id()
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT l.full_name, c.name AS class_name, l.room, l.phone
            FROM school_class_leaders l LEFT JOIN school_classes c ON c.id=l.class_id
            WHERE l.company_id=%s ORDER BY c.sort_order, c.name
        """, (company_id,))
        rows = cur.fetchall()
    finally:
        pool.putconn(conn)
    wb = Workbook(); ws = wb.active; ws.title = "Классные руководители"
    ws.append(["№", "ФИО", "Класс", "Кабинет", "Телефон"])
    for index, row in enumerate(rows, 1):
        ws.append([index, row["full_name"], row["class_name"] or "", row["room"] or "", row["phone"] or ""])
    _style_sheet(ws, [8, 38, 14, 14, 20])
    output = BytesIO(); wb.save(output); output.seek(0)
    return send_file(output, as_attachment=True, download_name="school_class_leaders.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@school_bp.route("/meals")
def meals():
    company_id = _company_id()
    if not company_id:
        return redirect("/login")
    try:
        selected_date = datetime.strptime(request.args.get("date", date.today().isoformat()), "%Y-%m-%d").date()
    except ValueError:
        selected_date = date.today()
    conn = get_db()
    try:
        cur = conn.cursor()
        prices = _prices_for_date(cur, company_id, selected_date)
        cur.execute("""
            SELECT c.id AS class_id, c.name AS class_name, l.full_name AS leader_name,
                   m.id, COALESCE(m.plan_count,0) plan_count, COALESCE(m.fact_count,0) fact_count,
                   COALESCE(m.free_count,0) free_count, COALESCE(m.paid_count,0) paid_count,
                   COALESCE(m.free_price,%s) free_price, COALESCE(m.paid_price,%s) paid_price,
                   COALESCE(m.note,'') note
            FROM school_classes c
            LEFT JOIN school_class_leaders l ON l.class_id=c.id AND l.company_id=c.company_id
            LEFT JOIN school_meals m ON m.class_id=c.id AND m.company_id=c.company_id AND m.meal_date=%s
            WHERE c.company_id=%s AND c.is_active=TRUE
            ORDER BY c.sort_order, c.name
        """, (prices["free_price"], prices["paid_price"], selected_date, company_id))
        rows = cur.fetchall()
    finally:
        pool.putconn(conn)
    totals = {
        "plan": sum(r["plan_count"] for r in rows), "fact": sum(r["fact_count"] for r in rows),
        "free": sum(r["free_count"] for r in rows), "paid": sum(r["paid_count"] for r in rows),
        "amount": sum((r["free_count"] * r["free_price"] + r["paid_count"] * r["paid_price"] for r in rows), Decimal("0")),
    }
    return render_template("school_meals.html", rows=rows, selected_date=selected_date, prices=prices, totals=totals)


@school_bp.route("/meals/save", methods=["POST"])
def save_meals():
    company_id = _company_id()
    meal_date = datetime.strptime(request.form.get("meal_date"), "%Y-%m-%d").date()
    class_ids = request.form.getlist("class_id")
    conn = get_db()
    try:
        cur = conn.cursor()
        default_prices = _prices_for_date(cur, company_id, meal_date)
        for class_id in class_ids:
            plan = _number(request.form.get(f"plan_{class_id}"))
            fact = _number(request.form.get(f"fact_{class_id}"))
            free = _number(request.form.get(f"free_{class_id}"))
            paid = _number(request.form.get(f"paid_{class_id}"))
            note = request.form.get(f"note_{class_id}", "").strip()
            if not any((plan, fact, free, paid, note)):
                cur.execute("DELETE FROM school_meals WHERE company_id=%s AND class_id=%s AND meal_date=%s", (company_id, class_id, meal_date))
                continue
            cur.execute("""
                INSERT INTO school_meals(company_id,class_id,meal_date,plan_count,fact_count,free_count,paid_count,free_price,paid_price,note,created_by)
                VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                ON CONFLICT(company_id,class_id,meal_date) DO UPDATE SET
                    plan_count=EXCLUDED.plan_count, fact_count=EXCLUDED.fact_count,
                    free_count=EXCLUDED.free_count, paid_count=EXCLUDED.paid_count,
                    note=EXCLUDED.note, updated_at=NOW()
            """, (company_id, class_id, meal_date, plan, fact, free, paid, default_prices["free_price"], default_prices["paid_price"], note, session.get("user_id")))
        conn.commit()
    except Exception:
        conn.rollback(); raise
    finally:
        pool.putconn(conn)
    return redirect(url_for("school.meals", date=meal_date.isoformat(), success="Данные сохранены"))


@school_bp.route("/meals/<int:meal_id>/delete", methods=["POST"])
def delete_meal(meal_id):
    company_id = _company_id()
    meal_date = request.form.get("meal_date") or date.today().isoformat()
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM school_meals WHERE id=%s AND company_id=%s", (meal_id, company_id))
        conn.commit()
    finally:
        pool.putconn(conn)
    return redirect(url_for("school.meals", date=meal_date, success="Строка удалена"))


@school_bp.route("/prices", methods=["POST"])
def save_prices():
    company_id = _company_id()
    effective_from = datetime.strptime(request.form.get("effective_from"), "%Y-%m-%d").date()
    free_price = _money(request.form.get("free_price")); paid_price = _money(request.form.get("paid_price"))
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO school_meal_prices(company_id,free_price,paid_price,effective_from)
            VALUES(%s,%s,%s,%s)
            ON CONFLICT(company_id,effective_from) DO UPDATE SET free_price=EXCLUDED.free_price, paid_price=EXCLUDED.paid_price
        """, (company_id, free_price, paid_price, effective_from))
        cur.execute("""
            UPDATE school_meals SET free_price=%s, paid_price=%s, updated_at=NOW()
            WHERE company_id=%s AND meal_date=%s
        """, (free_price, paid_price, company_id, effective_from))
        conn.commit()
    finally:
        pool.putconn(conn)
    return redirect(url_for("school.meals", date=effective_from.isoformat(), success="Цены сохранены"))


@school_bp.route("/meals/export.xlsx")
def export_meals():
    company_id = _company_id()
    date_from = request.args.get("date_from") or date.today().isoformat()
    date_to = request.args.get("date_to") or date_from
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT m.meal_date,c.name class_name,l.full_name leader_name,m.plan_count,m.fact_count,
                   m.free_count,m.paid_count,m.free_price,m.paid_price,m.note
            FROM school_meals m JOIN school_classes c ON c.id=m.class_id
            LEFT JOIN school_class_leaders l ON l.class_id=c.id AND l.company_id=c.company_id
            WHERE m.company_id=%s AND m.meal_date BETWEEN %s AND %s
            ORDER BY m.meal_date,c.sort_order,c.name
        """, (company_id, date_from, date_to))
        rows = cur.fetchall()
    finally:
        pool.putconn(conn)
    wb = Workbook(); ws = wb.active; ws.title = "Питание"
    ws.append(["Дата","Класс","Классный руководитель","План","Факт","Отсутствуют","Бесплатное","Платное","Всего питающихся","Цена бесплатного","Цена платного","Сумма","Примечание"])
    for row in rows:
        absent = max(0, row["plan_count"] - row["fact_count"])
        total = row["free_count"] + row["paid_count"]
        amount = row["free_count"] * row["free_price"] + row["paid_count"] * row["paid_price"]
        ws.append([row["meal_date"],row["class_name"],row["leader_name"] or "",row["plan_count"],row["fact_count"],absent,row["free_count"],row["paid_count"],total,row["free_price"],row["paid_price"],amount,row["note"] or ""])
    _style_sheet(ws, [14,12,34,10,10,14,14,12,18,20,18,16,30])
    for cell in ws["A"][1:]: cell.number_format = "yyyy-mm-dd"
    for col in ("J","K","L"):
        for cell in ws[col][1:]: cell.number_format = '#,##0.00'
    output=BytesIO(); wb.save(output); output.seek(0)
    return send_file(output,as_attachment=True,download_name=f"school_meals_{date_from}_{date_to}.xlsx",mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@school_bp.route("/meals/import", methods=["POST"])
def import_meals():
    company_id = _company_id(); upload = request.files.get("file")
    if not upload or not upload.filename.lower().endswith(".xlsx"):
        return redirect(url_for("school.meals", error="Выберите Excel-файл"))
    try:
        ws = load_workbook(upload, read_only=True, data_only=True).active
        headers = {_text(c.value).lower(): i for i, c in enumerate(ws[1])}
        required = {"дата", "класс", "план", "факт", "бесплатное", "платное"}
        if not required.issubset(headers):
            raise ValueError("Нужны колонки: Дата, Класс, План, Факт, Бесплатное, Платное")
        conn = get_db(); cur = conn.cursor(); count = 0
        try:
            for row in ws.iter_rows(min_row=2, values_only=True):
                raw_date = row[headers["дата"]]
                if not raw_date or not _text(row[headers["класс"]]): continue
                meal_date = raw_date.date() if isinstance(raw_date, datetime) else (raw_date if isinstance(raw_date, date) else datetime.strptime(_text(raw_date), "%Y-%m-%d").date())
                class_id = _ensure_class(cur, company_id, row[headers["класс"]])
                prices = _prices_for_date(cur, company_id, meal_date)
                note_col = headers.get("примечание")
                note = _text(row[note_col]) if note_col is not None else ""
                cur.execute("""
                    INSERT INTO school_meals(company_id,class_id,meal_date,plan_count,fact_count,free_count,paid_count,free_price,paid_price,note,created_by)
                    VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    ON CONFLICT(company_id,class_id,meal_date) DO UPDATE SET plan_count=EXCLUDED.plan_count,fact_count=EXCLUDED.fact_count,free_count=EXCLUDED.free_count,paid_count=EXCLUDED.paid_count,note=EXCLUDED.note,updated_at=NOW()
                """, (company_id,class_id,meal_date,_number(row[headers["план"]]),_number(row[headers["факт"]]),_number(row[headers["бесплатное"]]),_number(row[headers["платное"]]),prices["free_price"],prices["paid_price"],note,session.get("user_id")))
                count += 1
            conn.commit()
        except Exception:
            conn.rollback(); raise
        finally:
            pool.putconn(conn)
    except Exception as exc:
        return redirect(url_for("school.meals", error=f"Ошибка импорта: {exc}"))
    return redirect(url_for("school.meals", success=f"Загружено строк: {count}"))
